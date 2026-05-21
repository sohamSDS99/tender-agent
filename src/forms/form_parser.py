"""
Form Parser — Detects file format and extracts form fields/structure.

Supports:
- PDF: Fillable AcroForm fields, or text-based field extraction
- DOCX: Table-based forms, placeholder text, form content controls
- XLSX: Cell-based forms with labels and input cells
- Fallback: Extracts raw text for LLM-based field detection

Usage:
    parser = FormParser()
    result = parser.parse("/path/to/tender_form.pdf")
    print(result.fields)  # [FormField(...), ...]
"""

from __future__ import annotations

import os
import re
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any


class FileFormat(Enum):
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    UNKNOWN = "unknown"


@dataclass
class FormField:
    """A single field extracted from a form."""
    name: str                    # Field label / name
    field_type: str = "text"     # text, checkbox, radio, dropdown, date, number, textarea
    current_value: str = ""      # Pre-filled value if any
    page_or_section: str = ""    # Location context (page number, section name)
    required: bool = False       # Whether the field appears mandatory
    options: list[str] = field(default_factory=list)  # Dropdown/radio options
    metadata: dict[str, Any] = field(default_factory=dict)  # Format-specific info


@dataclass
class ParseResult:
    """Result of parsing a form file."""
    file_path: str
    file_format: FileFormat
    fields: list[FormField]
    raw_text: str = ""           # Full text content for LLM context
    page_count: int = 0
    title: str = ""
    parse_errors: list[str] = field(default_factory=list)


class FormParser:
    """Detects file format and extracts form fields."""

    def parse(self, file_path: str) -> ParseResult:
        """Parse a form file and extract its fields.

        Auto-detects format from file extension and content.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        fmt = self._detect_format(path)

        if fmt == FileFormat.PDF:
            return self._parse_pdf(path)
        elif fmt == FileFormat.DOCX:
            return self._parse_docx(path)
        elif fmt == FileFormat.XLSX:
            return self._parse_xlsx(path)
        else:
            return self._parse_text_fallback(path)

    def _detect_format(self, path: Path) -> FileFormat:
        """Detect file format from extension."""
        ext = path.suffix.lower()
        format_map = {
            ".pdf": FileFormat.PDF,
            ".docx": FileFormat.DOCX,
            ".doc": FileFormat.DOCX,
            ".xlsx": FileFormat.XLSX,
            ".xls": FileFormat.XLSX,
            ".csv": FileFormat.XLSX,
        }
        return format_map.get(ext, FileFormat.UNKNOWN)

    # -----------------------------------------------------------------------
    # PDF Parsing
    # -----------------------------------------------------------------------

    def _parse_pdf(self, path: Path) -> ParseResult:
        """Extract fields from a PDF form.

        Strategy, in order:

        1. AcroForm widgets — modern fillable PDFs. Each widget
           becomes a FormField with ``metadata['source']='acroform'``.
        2. GEOMETRIC extraction (flat-PDF inputs drawn as boxes) —
           uses the SAME box detection the writer uses, so every
           field emitted here carries a ``box_anchor`` the writer
           can place answers into without re-discovery. This replaces
           the brittle ``_extract_fields_from_text`` heuristic for
           any PDF that has visible input rectangles.
        3. Text-pattern fallback — only if both AcroForm AND
           geometric extraction find nothing. Last resort.
        """
        fields: list[FormField] = []
        raw_text = ""
        page_count = 0
        errors: list[str] = []

        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            page_count = len(reader.pages)

            # Strategy 1: AcroForm fields (fillable PDF)
            if reader.get_form_text_fields() is not None:
                form_fields = reader.get_form_text_fields()
                if form_fields:
                    for name, value in form_fields.items():
                        fields.append(FormField(
                            name=name,
                            field_type="text",
                            current_value=value or "",
                            metadata={"source": "acroform"},
                        ))

            # Also check for other field types (checkboxes, dropdowns)
            if reader.get_fields():
                for name, field_obj in reader.get_fields().items():
                    # Skip if already added as text field
                    if any(f.name == name for f in fields):
                        continue

                    field_type = field_obj.get("/FT", "")
                    ff = FormField(
                        name=name,
                        metadata={"source": "acroform", "pdf_type": str(field_type)},
                    )

                    if field_type == "/Btn":
                        ff.field_type = "checkbox"
                        value = field_obj.get("/V", "")
                        ff.current_value = "Yes" if value in ("/Yes", "/On", True) else ""
                    elif field_type == "/Ch":
                        ff.field_type = "dropdown"
                        opts = field_obj.get("/Opt", [])
                        ff.options = [str(o) for o in opts] if opts else []
                        ff.current_value = str(field_obj.get("/V", ""))
                    else:
                        ff.field_type = "text"
                        ff.current_value = str(field_obj.get("/V", "") or "")

                    fields.append(ff)

            # Always extract raw text for context
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                raw_text += f"\n--- Page {i + 1} ---\n{page_text}"

        except ImportError:
            errors.append("pypdf not installed — cannot parse PDF forms. Install with: pip install pypdf")
            # Try pdfplumber as fallback for text
            try:
                import pdfplumber
                with pdfplumber.open(str(path)) as pdf:
                    page_count = len(pdf.pages)
                    for i, page in enumerate(pdf.pages):
                        raw_text += f"\n--- Page {i + 1} ---\n{page.extract_text() or ''}"
            except ImportError:
                errors.append("pdfplumber not installed either — install pypdf or pdfplumber")
        except Exception as exc:
            errors.append(f"PDF parsing error: {exc}")

        # Strategy 2: GEOMETRIC extraction — used for flat PDFs that
        # have visible input rectangles drawn on the page but no real
        # AcroForm widgets. Reuses the writer's box detector so every
        # field carries a box_anchor and the writer can place answers
        # without re-detection or name-matching.
        if not fields:
            try:
                geom_fields, geom_errs = self._parse_pdf_geometric(path)
                fields.extend(geom_fields)
                errors.extend(geom_errs)
            except Exception as exc:  # noqa: BLE001
                # Geometric parser SHOULD be robust; if it crashes,
                # record and fall through to text-pattern fallback.
                errors.append(f"Geometric extraction failed: {exc}")

        # Strategy 3: text-pattern fallback. Only if nothing else
        # produced fields. This path produced lots of garbage in the
        # past (e.g. "iod" from "Period") so we only hit it when both
        # AcroForm and geometry came up empty.
        if not fields and raw_text:
            fields = self._extract_fields_from_text(raw_text)

        return ParseResult(
            file_path=str(path),
            file_format=FileFormat.PDF,
            fields=fields,
            raw_text=raw_text.strip(),
            page_count=page_count,
            title=path.stem,
            parse_errors=errors,
        )

    def _parse_pdf_geometric(self, path: Path) -> tuple[list[FormField], list[str]]:
        """Extract one field per detected input rectangle on a flat PDF.

        Returns (fields, errors). The fields carry their box geometry
        on ``metadata['box_anchor']`` so the writer can place answers
        into the exact same rectangle — no re-detection, no label
        matching against the field name.

        Labels are deduplicated by appending positional disambiguators
        (page + ordinal) when the same label appears more than once
        (e.g. budget tables with many "$" amount boxes).
        """
        from collections import Counter
        from .pdf_geometry import (
            BoxAnchor, build_label_context, detect_input_boxes,
        )

        try:
            import pdfplumber
        except ImportError:
            return [], ["pdfplumber not installed — cannot do geometric extraction"]

        fields: list[FormField] = []
        errors: list[str] = []

        with pdfplumber.open(str(path)) as pdf:
            # First pass: collect (page_idx, box, label) for every detected box.
            raw: list[tuple[int, dict, str, float, float]] = []
            for page_idx, page in enumerate(pdf.pages):
                page_w, page_h = float(page.width), float(page.height)
                boxes = detect_input_boxes(page)
                words = page.extract_words(
                    x_tolerance=2, y_tolerance=2,
                    keep_blank_chars=False, use_text_flow=True,
                )
                for b in boxes:
                    label = build_label_context(b, words).strip()
                    raw.append((page_idx, b, label, page_w, page_h))

            if not raw:
                return [], []

            # Disambiguate duplicate labels. If "$" appears on 80
            # boxes, we'd lose every box but one in the FormFiller's
            # name-keyed dict. Append "(p.N #M)" where N is the page
            # and M is the within-page occurrence index of the label.
            per_page_label_count: dict[tuple[int, str], int] = {}
            label_total = Counter(label for _, _, label, _, _ in raw if label)
            for page_idx, b, label, page_w, page_h in raw:
                effective = label or f"unlabeled field p.{page_idx + 1}"

                # If this label is unique across the whole doc, leave
                # it as-is — gives the LLM a clean, unambiguous name.
                if label_total.get(label, 0) > 1 or not label:
                    key = (page_idx, effective)
                    per_page_label_count[key] = per_page_label_count.get(key, 0) + 1
                    occ = per_page_label_count[key]
                    effective = f"{effective} (p.{page_idx + 1} #{occ})"

                anchor = BoxAnchor(
                    page_idx=page_idx,
                    x0=float(b["x0"]), top=float(b["top"]),
                    x1=float(b["x1"]), bottom=float(b["bottom"]),
                    page_width=page_w, page_height=page_h,
                )
                fields.append(FormField(
                    name=effective,
                    field_type="text",
                    page_or_section=f"page {page_idx + 1}",
                    metadata={
                        "source": "geometric",
                        "box_anchor": anchor.to_dict(),
                        "raw_label": label,
                    },
                ))

        return fields, errors

    # -----------------------------------------------------------------------
    # DOCX Parsing
    # -----------------------------------------------------------------------

    def _parse_docx(self, path: Path) -> ParseResult:
        """Extract fields from a Word document form.

        Strategy:
        1. Look for table-based forms (label | input pattern)
        2. Look for placeholder text like [Company Name], {{field}}, ____
        3. Look for content controls (structured document tags)
        """
        fields: list[FormField] = []
        raw_text = ""
        errors: list[str] = []

        try:
            from docx import Document

            doc = Document(str(path))

            # Extract all text for context
            paragraphs = []
            for para in doc.paragraphs:
                paragraphs.append(para.text)
            raw_text = "\n".join(paragraphs)

            # Strategy 1: Table-based forms
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]

                    if len(cells) >= 2:
                        # Pattern: Label | Value (empty or placeholder)
                        label = cells[0]
                        value = cells[1] if len(cells) > 1 else ""

                        # Skip header rows and rows where both cells are short labels
                        if label and len(label) > 1 and not label.startswith("---"):
                            is_fillable = (
                                not value
                                or value.startswith("[")
                                or value.startswith("{")
                                or "_" * 3 in value
                                or value.lower() in ("n/a", "tbd", "enter", "please enter")
                            )
                            if is_fillable or (label.endswith(":") or label.endswith("?")):
                                ff = FormField(
                                    name=label.rstrip(":").rstrip("?").strip(),
                                    current_value="" if is_fillable else value,
                                    page_or_section=f"Table {table_idx + 1}, Row {row_idx + 1}",
                                    metadata={"source": "docx_table"},
                                )
                                # Avoid duplicate field names
                                if not any(f.name == ff.name for f in fields):
                                    fields.append(ff)

            # Strategy 2: Placeholder patterns in paragraphs
            placeholder_patterns = [
                r'\[([^\]]{2,60})\]',           # [Company Name]
                r'\{\{([^}]{2,60})\}\}',         # {{company_name}}
                r'_{3,}',                          # _________ (blank line)
                r'<([A-Z][^>]{2,40})>',          # <COMPANY NAME>
            ]
            for para in doc.paragraphs:
                text = para.text.strip()
                for pattern in placeholder_patterns:
                    for match in re.finditer(pattern, text):
                        if match.groups():
                            name = match.group(1).strip()
                        else:
                            # For ___ patterns, try to extract label from context
                            before = text[:match.start()].strip().rstrip(":")
                            name = before if before else "Blank Field"

                        if name and len(name) > 1:
                            ff = FormField(
                                name=name,
                                metadata={"source": "docx_placeholder", "context": text[:100]},
                            )
                            if not any(f.name == ff.name for f in fields):
                                fields.append(ff)

        except ImportError:
            errors.append("python-docx not installed — install with: pip install python-docx")
        except Exception as exc:
            errors.append(f"DOCX parsing error: {exc}")

        # Fallback: extract from raw text
        if not fields and raw_text:
            fields = self._extract_fields_from_text(raw_text)

        return ParseResult(
            file_path=str(path),
            file_format=FileFormat.DOCX,
            fields=fields,
            raw_text=raw_text.strip(),
            title=path.stem,
            parse_errors=errors,
        )

    # -----------------------------------------------------------------------
    # XLSX Parsing
    # -----------------------------------------------------------------------

    def _parse_xlsx(self, path: Path) -> ParseResult:
        """Extract fields from an Excel form.

        Strategy: Look for label-value pairs where label is in column A/B
        and value cell is empty or contains a placeholder.
        """
        fields: list[FormField] = []
        raw_text = ""
        errors: list[str] = []

        try:
            from openpyxl import load_workbook

            wb = load_workbook(str(path), data_only=True)

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                raw_text += f"\n--- Sheet: {sheet_name} ---\n"

                for row in ws.iter_rows(min_row=1, max_row=ws.max_row, values_only=False):
                    row_texts = []
                    for cell in row:
                        val = str(cell.value) if cell.value is not None else ""
                        row_texts.append(val)

                    raw_text += " | ".join(row_texts) + "\n"

                    # Look for label-value pairs
                    for i in range(len(row) - 1):
                        label_cell = row[i]
                        value_cell = row[i + 1]

                        label = str(label_cell.value or "").strip()
                        value = str(value_cell.value or "").strip()

                        if not label or len(label) < 2:
                            continue

                        # Check if this looks like a label → input pair
                        is_label = (
                            label.endswith(":")
                            or label.endswith("?")
                            or label[0].isupper()
                        )
                        is_empty_value = (
                            not value
                            or value.startswith("[")
                            or value.startswith("{")
                            or value.lower() in ("n/a", "tbd", "enter here")
                        )

                        if is_label and is_empty_value:
                            ff = FormField(
                                name=label.rstrip(":").rstrip("?").strip(),
                                current_value="" if is_empty_value else value,
                                page_or_section=f"Sheet: {sheet_name}, Row {label_cell.row}",
                                metadata={
                                    "source": "xlsx",
                                    "cell_ref": f"{value_cell.column_letter}{value_cell.row}",
                                    "sheet": sheet_name,
                                },
                            )
                            if not any(f.name == ff.name for f in fields):
                                fields.append(ff)

        except ImportError:
            errors.append("openpyxl not installed — install with: pip install openpyxl")
        except Exception as exc:
            errors.append(f"XLSX parsing error: {exc}")

        if not fields and raw_text:
            fields = self._extract_fields_from_text(raw_text)

        return ParseResult(
            file_path=str(path),
            file_format=FileFormat.XLSX,
            fields=fields,
            raw_text=raw_text.strip(),
            title=path.stem,
            parse_errors=errors,
        )

    # -----------------------------------------------------------------------
    # Text-based fallback field extraction
    # -----------------------------------------------------------------------

    def _parse_text_fallback(self, path: Path) -> ParseResult:
        """For unknown formats, read as text and extract fields."""
        try:
            raw_text = path.read_text(encoding="utf-8", errors="replace")
        except Exception:
            raw_text = ""

        fields = self._extract_fields_from_text(raw_text) if raw_text else []

        return ParseResult(
            file_path=str(path),
            file_format=FileFormat.UNKNOWN,
            fields=fields,
            raw_text=raw_text,
            title=path.stem,
        )

    def _extract_fields_from_text(self, text: str) -> list[FormField]:
        """Extract form fields from raw text using pattern matching.

        Looks for common form patterns:
        - "Label: _____" or "Label: "
        - "Label .............. "
        - Lines ending with ":"
        """
        fields: list[FormField] = []
        seen_names: set[str] = set()

        # Pattern 1: "Label: ___" or "Label: " with trailing whitespace
        for match in re.finditer(r'^(.{3,60}):\s*[_.\s]{3,}', text, re.MULTILINE):
            name = match.group(1).strip()
            if name and name not in seen_names:
                fields.append(FormField(name=name, metadata={"source": "text_pattern"}))
                seen_names.add(name)

        # Pattern 2: Lines that look like form labels (end with colon, followed by empty/short content)
        for match in re.finditer(r'^(.{3,60}):\s*$', text, re.MULTILINE):
            name = match.group(1).strip()
            if name and name not in seen_names and not name.startswith("#"):
                fields.append(FormField(name=name, metadata={"source": "text_colon"}))
                seen_names.add(name)

        return fields
