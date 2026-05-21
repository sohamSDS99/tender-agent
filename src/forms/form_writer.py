"""
Form Writer — Writes filled values back into the original file format.

Supports:
- PDF: Fill AcroForm fields, or create annotated overlay
- DOCX: Replace placeholders, fill table cells
- XLSX: Write values to identified cells

Usage:
    writer = FormWriter()
    output_path = writer.write(
        original_path="/path/to/form.pdf",
        filled_fields=[FilledField(name="Company", value="SDS Manager", ...)],
        output_path="/path/to/filled_form.pdf",
    )
"""

from __future__ import annotations

import os
import re
import shutil
from pathlib import Path
from typing import Any

from .form_parser import FileFormat, FormField
from .form_filler import FilledField


def _numeric_to_tier(confidence: float) -> str:
    """Fallback tier mapping when FilledField.confidence_tier is unset.

    Mirrors the thresholds in form_filler._tier_for so the canonical
    exports show the same colour-coding as the Approvals UI.
    """
    try:
        c = float(confidence)
    except (TypeError, ValueError):
        return "low"
    if c >= 0.8:
        return "high"
    if c >= 0.55:
        return "medium"
    return "low"


def _escape_pdf(text: str) -> str:
    """Escape characters that would otherwise break reportlab Paragraph parsing."""
    if not text:
        return ""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


class FormWriter:
    """Writes filled values back into form files."""

    def write(
        self,
        original_path: str,
        filled_fields: list[FilledField],
        output_path: str | None = None,
    ) -> str:
        """Write filled values into a copy of the original form.

        Args:
            original_path: Path to the original form file.
            filled_fields: List of fields with values to write.
            output_path: Where to save the filled form. Defaults to
                         original_name_FILLED.ext in the same directory.

        Returns:
            Path to the filled form file.
        """
        path = Path(original_path)
        if not path.exists():
            raise FileNotFoundError(f"Original file not found: {original_path}")

        # Determine output path
        if not output_path:
            stem = path.stem
            output_path = str(path.parent / f"{stem}_FILLED{path.suffix}")

        # Detect format and dispatch
        ext = path.suffix.lower()
        if ext == ".pdf":
            return self._write_pdf(path, filled_fields, output_path)
        elif ext in (".docx", ".doc"):
            return self._write_docx(path, filled_fields, output_path)
        elif ext in (".xlsx", ".xls"):
            return self._write_xlsx(path, filled_fields, output_path)
        else:
            return self._write_text_fallback(path, filled_fields, output_path)

    # -----------------------------------------------------------------------
    # PDF Writing
    # -----------------------------------------------------------------------

    def _write_pdf(
        self, path: Path, fields: list[FilledField], output_path: str
    ) -> str:
        """Fill a PDF form.

        Two-strategy dispatch:

          A. If the PDF has real AcroForm widget annotations (most
             modern federal forms), fill them via pypdf's
             update_page_form_field_values + set NeedAppearances=true
             so the appearance streams regenerate on open.

          B. If the PDF is "flat" (visual input rectangles drawn on
             the page, no real form widgets — common for older
             federal templates like the HHS Subcontracting Plan), use
             a pdfplumber + reportlab + pypdf overlay pipeline that
             searches for each field's label by text and draws the
             answer at the right (x, y) position.  See
             ``_write_pdf_overlay`` for the algorithm.

        Detection runs first and is cheap — the same pdfplumber/
        reportlab/pypdf install supports both code paths.
        """
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
        except ImportError:
            shutil.copy2(str(path), output_path)
            print(f"pypdf not available — copied original to {output_path}")
            return output_path
        except Exception as exc:
            print(f"PDF read error: {exc}")
            shutil.copy2(str(path), output_path)
            return output_path

        # AcroForm presence check — get_fields() returns None or {} for
        # flat PDFs and a populated dict for real fillable forms.
        try:
            acroform_fields = reader.get_fields() or {}
        except Exception as exc:  # noqa: BLE001
            print(f"  [pdf-fill] AcroForm probe failed ({exc}) — assuming flat PDF")
            acroform_fields = {}

        if acroform_fields:
            return self._write_pdf_acroform_fill(
                reader, path, fields, output_path
            )
        else:
            print(
                "  [pdf-fill] PDF has no AcroForm widgets — "
                "falling back to text-overlay strategy."
            )
            return self._write_pdf_overlay_fill(path, fields, output_path)

    # -----------------------------------------------------------------------
    # PDF Strategy A: AcroForm widget fill (modern fillable PDFs)
    # -----------------------------------------------------------------------

    def _write_pdf_acroform_fill(
        self,
        reader: "pypdf.PdfReader",  # noqa: F821 — pypdf already imported above
        path: Path,
        fields: list[FilledField],
        output_path: str,
    ) -> str:
        """Strategy A — fill AcroForm widget values.

        Two historical bugs that used to make this path output an
        "empty-looking" PDF, both fixed here:

          1. ``PdfWriter()`` + per-page ``add_page()`` copies the page
             content but NOT the document catalog's ``/AcroForm``
             dictionary, so readers don't render widgets as fields.
             Solution: ``PdfWriter(clone_from=reader)``.

          2. Setting ``/V`` doesn't regenerate the cached appearance
             stream, so the field text isn't visible even though the
             data is stored.  Solution: set
             ``/AcroForm /NeedAppearances = true``.
        """
        import pypdf
        from pypdf.generic import NameObject, BooleanObject

        field_map = {f.name: f.value for f in fields if f.value}

        try:
            try:
                writer = pypdf.PdfWriter(clone_from=reader)
            except TypeError:
                # Older pypdf versions — graft AcroForm manually.
                writer = pypdf.PdfWriter()
                for page in reader.pages:
                    writer.add_page(page)
                try:
                    src_root = reader.trailer["/Root"]
                    if "/AcroForm" in src_root:
                        writer._root_object[NameObject("/AcroForm")] = (
                            src_root["/AcroForm"]
                        )
                except Exception as exc:  # noqa: BLE001
                    print(f"  [pdf-fill] AcroForm graft failed: {exc}")

            if "/AcroForm" in writer._root_object:
                writer._root_object["/AcroForm"][
                    NameObject("/NeedAppearances")
                ] = BooleanObject(True)

            matched_total = 0
            if field_map:
                for i, page in enumerate(writer.pages):
                    try:
                        writer.update_page_form_field_values(page, field_map)
                        if "/Annots" in page:
                            for annot_ref in page["/Annots"]:
                                annot = annot_ref.get_object()
                                if (
                                    annot.get("/Subtype") == "/Widget"
                                    and annot.get("/T") in field_map
                                ):
                                    matched_total += 1
                    except Exception as exc:  # noqa: BLE001
                        print(
                            f"  [pdf-fill] update_page_form_field_values "
                            f"failed on page {i}: {exc}"
                        )

            with open(output_path, "wb") as f:
                writer.write(f)

            print(
                f"  [pdf-fill] AcroForm fill → {output_path} — "
                f"matched ~{matched_total}/{len(field_map)} field(s)."
            )
            return output_path

        except Exception as exc:
            print(f"  [pdf-fill] AcroForm path crashed: {exc}")
            shutil.copy2(str(path), output_path)
            return output_path

    # -----------------------------------------------------------------------
    # PDF Strategy B: Text overlay for flat PDFs (no AcroForm widgets)
    # -----------------------------------------------------------------------

    def _write_pdf_overlay_fill(
        self,
        path: Path,
        fields: list[FilledField],
        output_path: str,
    ) -> str:
        """Strategy B — geometry-driven overlay for flat PDFs.

        Permanent design (does NOT use the brittle "text-position
        next to label" heuristic from an earlier draft).  This version
        works on ANY flat-PDF form because it locates the actual
        drawn input rectangles and writes inside them:

          1. Detect every input rectangle on each page by pairing
             thin horizontal edge-strips (top + bottom) with matching
             x range, 8-30pt apart vertically.  These are the boxes
             the user sees as visual input areas.
          2. Dedupe boxes that overlap >70% by area — many PDFs draw
             each input as a frame-within-a-frame (inner + outer
             border), producing two overlapping detections.  Keep
             the outer one.
          3. For each box, build a "label context" by collecting
             every word in the rectangle:
                  x: [box.x0 - LABEL_LEFT_REACH, box.x0]
                  y: [box.top - LABEL_ABOVE_REACH, box.bottom + 2]
             ...joined into a normalised string.  This captures
             multi-line wrapped labels (e.g. "Total Contract
             Amount (including options, and any modifications if
             this submission is due to a modification):") AND
             same-row labels in a single search space.
          4. For each filled field, find the BEST-matching unused
             box whose label context contains the field name as a
             substring (after normalisation).  Longer field names
             matched first so "Title of Acquisition" claims its box
             before bare "Title" does.
          5. Place the answer text INSIDE the matched box, with 4pt
             padding from the left edge, vertically baselined at
             box.bottom - 3pt.  Font size derived from box height,
             clamped to 7..12pt.  Long answers are truncated to fit.
          6. Generate overlay PDF with reportlab; composite with
             pypdf via merge_page.

        Misalignment guards:
            - Text is constrained to live INSIDE the detected box —
              cannot collide with adjacent label text or other fields.
            - Each box is used at most once.
            - Longer field names matched first → no shorter-name
              false matches inside longer labels.
            - Box width determines max_width → no horizontal overflow.
            - Single per-field log line per placement + per-miss so
              you can audit alignment from the bridge log.
        """
        try:
            import pdfplumber
            import pypdf
            from io import BytesIO
            from reportlab.pdfgen import canvas
        except ImportError as exc:
            print(f"  [pdf-overlay] Required lib missing: {exc} — copying blank.")
            shutil.copy2(str(path), output_path)
            return output_path

        field_map = {f.name: f.value for f in fields if f.value}
        if not field_map:
            shutil.copy2(str(path), output_path)
            print(f"  [pdf-overlay] No values to overlay — copied blank.")
            return output_path

        # ------------------------------------------------------------------
        # 1) Parse PDF: collect page geometry + per-box label context.
        # ------------------------------------------------------------------
        try:
            page_payloads: list[dict] = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    page_w, page_h = float(page.width), float(page.height)
                    boxes = _detect_input_boxes(page)
                    words = page.extract_words(
                        x_tolerance=2, y_tolerance=2,
                        keep_blank_chars=False, use_text_flow=True,
                    )
                    # Compute a label context per box.
                    for b in boxes:
                        b["label_context"] = _build_label_context(b, words)
                    page_payloads.append({
                        "width": page_w,
                        "height": page_h,
                        "boxes": boxes,
                    })
        except Exception as exc:
            print(f"  [pdf-overlay] pdfplumber parse failed: {exc} — copying blank.")
            shutil.copy2(str(path), output_path)
            return output_path

        total_boxes = sum(len(p["boxes"]) for p in page_payloads)
        print(
            f"  [pdf-overlay] Detected {total_boxes} input rectangle(s) across "
            f"{len(page_payloads)} page(s)."
        )

        # ------------------------------------------------------------------
        # 2) Match each field to a detected box (longest-name first).
        # ------------------------------------------------------------------
        placements: list[dict] = []
        misses: list[str] = []
        consumed_box_ids: set[tuple[int, int]] = set()  # (page_idx, box_idx)

        sorted_fields = sorted(
            ((k, v) for k, v in field_map.items() if v),
            key=lambda kv: (-len(_normalise_label(kv[0]).split()), kv[0]),
        )

        for field_name, value in sorted_fields:
            target = _normalise_label(field_name)
            if not target:
                misses.append(field_name)
                continue

            best: tuple[int, int, dict] | None = None
            best_score = 0.0
            for page_idx, payload in enumerate(page_payloads):
                for box_idx, box in enumerate(payload["boxes"]):
                    if (page_idx, box_idx) in consumed_box_ids:
                        continue
                    context = box.get("label_context", "")
                    if not context or target not in context:
                        continue
                    # Specificity score = how much of the box's label
                    # context the field name covers.
                    #   ctx = "title of acquisition", target = "title of
                    #         acquisition" → score = 1.0 (perfect fit)
                    #   ctx = "option 3 (if applicable): $ $",
                    #     target = "option 3"             → score = 0.29
                    #   ctx = "(if applicable) $ option 1 (if applicable):
                    #          $ option 3 (if applicable)",
                    #     target = "option 3"             → score = 0.12
                    # Higher = the field name dominates the label
                    # context = more confident it's the right box.
                    # This naturally prefers clean Strategy-1 hits
                    # over Strategy-2 mish-mashes that pull in text
                    # from neighbouring fields.
                    score = len(target) / max(1, len(context))
                    if score > best_score:
                        best_score = score
                        best = (page_idx, box_idx, box)

            if best is None:
                misses.append(field_name)
                continue

            page_idx, box_idx, box = best
            consumed_box_ids.add((page_idx, box_idx))

            # Position inside the box.
            x = float(box["x0"]) + 4.0  # 4pt left padding
            # PDF baseline is at the BOTTOM of glyphs.  We want them
            # vertically centred in the box.  Helvetica's baseline is
            # ~80% from the top of its bounding box, so baseline ≈
            # box.bottom - 3pt looks centred for ~10pt fonts.
            y_baseline_pdf = float(box["bottom"]) - 3.0
            font_size = max(7.0, min(12.0, float(box["height"]) - 4.0))
            max_width = max(20.0, float(box["x1"]) - float(box["x0"]) - 8.0)
            text = _truncate_to_width(str(value), max_width, font_size)

            placements.append({
                "page_idx": page_idx,
                "x": x,
                "y_baseline_pdf": y_baseline_pdf,
                "font_size": font_size,
                "text": text,
                "field_name": field_name,
                "box": box,
            })

        # ------------------------------------------------------------------
        # 3) Build the overlay PDF (transparent → original PDF below).
        #    pdfplumber: y origin top-left; reportlab: y origin bottom-left.
        #    Convert y_reportlab = page_height - y_pdfplumber.
        # ------------------------------------------------------------------
        overlay_buf = BytesIO()
        c = canvas.Canvas(overlay_buf)
        for page_idx, payload in enumerate(page_payloads):
            pw, ph = payload["width"], payload["height"]
            c.setPageSize((pw, ph))
            for p in placements:
                if p["page_idx"] != page_idx:
                    continue
                c.setFont("Helvetica", p["font_size"])
                c.setFillColorRGB(0, 0, 0)
                y_reportlab = ph - p["y_baseline_pdf"]
                c.drawString(p["x"], y_reportlab, p["text"])
            c.showPage()
        c.save()
        overlay_buf.seek(0)

        # ------------------------------------------------------------------
        # 4) Composite overlay onto the original PDF page-for-page.
        # ------------------------------------------------------------------
        try:
            base_reader = pypdf.PdfReader(str(path))
            overlay_reader = pypdf.PdfReader(overlay_buf)
            writer = pypdf.PdfWriter()
            for i, page in enumerate(base_reader.pages):
                if i < len(overlay_reader.pages):
                    page.merge_page(overlay_reader.pages[i])
                writer.add_page(page)
            with open(output_path, "wb") as f:
                writer.write(f)
        except Exception as exc:
            print(f"  [pdf-overlay] composite failed: {exc} — copying blank.")
            shutil.copy2(str(path), output_path)
            return output_path

        print(
            f"  [pdf-overlay] Wrote {output_path} — "
            f"placed {len(placements)}/{len(field_map)} answer(s) inside detected boxes."
        )
        for p in placements:
            b = p["box"]
            print(
                f"    ✓ p.{p['page_idx']+1} box[x={b['x0']:.0f}..{b['x1']:.0f}, "
                f"y={b['top']:.0f}..{b['bottom']:.0f}]  "
                f"font={p['font_size']:.1f}pt  '{p['field_name'][:40]}' "
                f"→ {p['text'][:30]!r}"
            )
        if misses:
            print(
                f"  [pdf-overlay] {len(misses)} field(s) had no box match: "
                f"{[m[:40] for m in misses[:6]]}{'…' if len(misses) > 6 else ''}"
            )
        return output_path

    # -----------------------------------------------------------------------
    # DOCX Writing
    # -----------------------------------------------------------------------

    def _write_docx(
        self, path: Path, fields: list[FilledField], output_path: str
    ) -> str:
        """Fill a Word document form.

        Strategy 1: Replace values in table cells
        Strategy 2: Replace placeholder patterns in paragraphs
        """
        field_map = {f.name: f.value for f in fields if f.value}

        try:
            from docx import Document

            doc = Document(str(path))

            # Strategy 1: Fill table cells
            for table in doc.tables:
                for row in table.rows:
                    cells = row.cells
                    if len(cells) >= 2:
                        label = cells[0].text.strip().rstrip(":").rstrip("?").strip()
                        if label in field_map:
                            cells[1].text = field_map[label]

            # Strategy 2: Replace placeholder patterns
            for para in doc.paragraphs:
                original_text = para.text

                for name, value in field_map.items():
                    # Replace [Field Name] → value
                    pattern1 = f"[{name}]"
                    if pattern1 in original_text:
                        original_text = original_text.replace(pattern1, value)

                    # Replace {{field_name}} → value
                    snake_name = name.lower().replace(" ", "_")
                    pattern2 = f"{{{{{snake_name}}}}}"
                    if pattern2 in original_text:
                        original_text = original_text.replace(pattern2, value)

                    # Replace <FIELD NAME> → value
                    pattern3 = f"<{name.upper()}>"
                    if pattern3 in original_text:
                        original_text = original_text.replace(pattern3, value)

                # Replace _____ blanks near matching labels
                for name, value in field_map.items():
                    label_pattern = re.escape(name) + r'\s*:?\s*[_]{3,}'
                    replacement = f"{name}: {value}"
                    original_text = re.sub(label_pattern, replacement, original_text, flags=re.IGNORECASE)

                if original_text != para.text:
                    # Clear existing runs and set new text
                    # Preserve formatting of the first run
                    if para.runs:
                        first_run = para.runs[0]
                        for run in para.runs[1:]:
                            run.text = ""
                        first_run.text = original_text
                    else:
                        para.text = original_text

            doc.save(output_path)
            print(f"DOCX form filled: {output_path}")
            return output_path

        except ImportError:
            shutil.copy2(str(path), output_path)
            print(f"python-docx not available — copied original to {output_path}")
            return output_path

        except Exception as exc:
            print(f"DOCX write error: {exc}")
            shutil.copy2(str(path), output_path)
            return output_path

    # -----------------------------------------------------------------------
    # XLSX Writing
    # -----------------------------------------------------------------------

    def _write_xlsx(
        self, path: Path, fields: list[FilledField], output_path: str
    ) -> str:
        """Fill an Excel form by writing to identified cells."""
        field_map = {f.name: f for f in fields if f.value}

        try:
            from openpyxl import load_workbook

            wb = load_workbook(str(path))

            for ff in fields:
                if not ff.value or not ff.original_field:
                    continue

                meta = ff.original_field.metadata
                cell_ref = meta.get("cell_ref", "")
                sheet_name = meta.get("sheet", "")

                if cell_ref and sheet_name and sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    ws[cell_ref] = ff.value
                else:
                    # Fallback: search for the label in all sheets and fill adjacent cell
                    for ws in wb.worksheets:
                        for row in ws.iter_rows():
                            for i, cell in enumerate(row):
                                label = str(cell.value or "").strip().rstrip(":").rstrip("?").strip()
                                if label == ff.name and i + 1 < len(row):
                                    row[i + 1].value = ff.value
                                    break

            wb.save(output_path)
            print(f"XLSX form filled: {output_path}")
            return output_path

        except ImportError:
            shutil.copy2(str(path), output_path)
            print(f"openpyxl not available — copied original to {output_path}")
            return output_path

        except Exception as exc:
            print(f"XLSX write error: {exc}")
            shutil.copy2(str(path), output_path)
            return output_path

    # -----------------------------------------------------------------------
    # Canonical multi-format export
    # -----------------------------------------------------------------------

    def write_canonical_set(
        self,
        filled_fields: list[FilledField],
        output_dir: str,
        base_name: str,
        *,
        title: str = "Filled Tender Form",
        include_low_confidence: bool = True,
    ) -> dict[str, str]:
        """Render the filled answers into three CLEAN editable formats.

        Unlike `write()` (which preserves the original document's
        template, logos, sections, etc.), this method generates
        canonical answer-only documents — flat Q&A tables ready for
        review, editing, or sharing.

        Always returns paths to all three formats; if a renderer
        fails the path is still returned but the file is best-effort.

        Args:
            filled_fields: The answers produced by FormFiller.
            output_dir: Directory to write the files into.
            base_name: Filename stem (no extension). The output files
                       will be `<base_name>_answers.{docx,xlsx,pdf}`.
            title: Human-readable title rendered at the top of each doc.
            include_low_confidence: When True, low-confidence fields are
                       rendered with a marker so the reviewer notices.

        Returns:
            Dict mapping format → absolute output path. Always has
            keys "docx", "xlsx", "pdf".
        """
        os.makedirs(output_dir, exist_ok=True)
        docx_path = os.path.join(output_dir, f"{base_name}_answers.docx")
        xlsx_path = os.path.join(output_dir, f"{base_name}_answers.xlsx")
        pdf_path = os.path.join(output_dir, f"{base_name}_answers.pdf")

        rows = self._canonical_rows(filled_fields, include_low_confidence)

        try:
            self._write_canonical_docx(rows, docx_path, title=title)
        except Exception as exc:
            print(f"[canonical docx] {type(exc).__name__}: {exc}")

        try:
            self._write_canonical_xlsx(rows, xlsx_path, title=title)
        except Exception as exc:
            print(f"[canonical xlsx] {type(exc).__name__}: {exc}")

        try:
            self._write_canonical_pdf(rows, pdf_path, title=title)
        except Exception as exc:
            print(f"[canonical pdf] {type(exc).__name__}: {exc}")

        return {"docx": docx_path, "xlsx": xlsx_path, "pdf": pdf_path}

    # ---- canonical row builder ----

    @staticmethod
    def _canonical_rows(
        fields: list[FilledField],
        include_low_confidence: bool,
    ) -> list[dict[str, str]]:
        """Flatten FilledField -> {question, answer, confidence_tier}."""
        rows: list[dict[str, str]] = []
        for ff in fields:
            value = (ff.value or "").strip()
            if not value:
                if not include_low_confidence:
                    continue
                value = "(needs human input)"
            tier = getattr(ff, "confidence_tier", "") or ""
            rows.append({
                "question": (ff.name or "").strip(),
                "answer": value,
                "confidence": tier or _numeric_to_tier(ff.confidence),
                "source": (ff.source or "").strip(),
            })
        return rows

    # ---- DOCX (editable) ----

    @staticmethod
    def _write_canonical_docx(
        rows: list[dict[str, str]],
        output_path: str,
        *,
        title: str,
    ) -> None:
        """Render the answers as a clean, editable DOCX.

        Layout: title at top, summary line, then a 3-column table
        (Question | Answer | Confidence). The DOCX is fully editable
        in Word, Google Docs, LibreOffice.
        """
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()

        # --- Title ---
        heading = doc.add_heading(title, level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

        # --- Summary line ---
        high = sum(1 for r in rows if r["confidence"] == "high")
        medium = sum(1 for r in rows if r["confidence"] == "medium")
        low = sum(1 for r in rows if r["confidence"] == "low")
        summary = doc.add_paragraph()
        summary_run = summary.add_run(
            f"{len(rows)} answers · {high} high · {medium} medium · {low} low confidence"
        )
        summary_run.italic = True
        summary_run.font.size = Pt(10)
        summary_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        # --- Q&A table ---
        if rows:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False

            header_cells = table.rows[0].cells
            for idx, label in enumerate(("Question", "Answer", "Confidence")):
                cell = header_cells[idx]
                cell.text = label
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(10)

            # Column widths
            try:
                table.columns[0].width = Inches(2.4)
                table.columns[1].width = Inches(3.8)
                table.columns[2].width = Inches(1.0)
            except Exception:
                pass

            for r in rows:
                cells = table.add_row().cells
                cells[0].text = r["question"]
                cells[1].text = r["answer"]
                cells[2].text = (r["confidence"] or "—").upper()
                # subtle styling: dim low-confidence cells
                if r["confidence"] == "low":
                    for cell in cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

        doc.save(output_path)

    # ---- XLSX (editable) ----

    @staticmethod
    def _write_canonical_xlsx(
        rows: list[dict[str, str]],
        output_path: str,
        *,
        title: str,
    ) -> None:
        """Render the answers as a clean, editable XLSX spreadsheet."""
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        ws = wb.active
        ws.title = "Answers"

        # Title row
        ws["A1"] = title
        ws["A1"].font = Font(size=14, bold=True, color="1F2937")
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=3)

        # Summary row
        high = sum(1 for r in rows if r["confidence"] == "high")
        medium = sum(1 for r in rows if r["confidence"] == "medium")
        low = sum(1 for r in rows if r["confidence"] == "low")
        ws["A2"] = (
            f"{len(rows)} answers · {high} high · {medium} medium · {low} low confidence"
        )
        ws["A2"].font = Font(italic=True, color="6B7280", size=10)
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=3)

        # Header row
        headers = ["Question", "Answer", "Confidence"]
        for col_idx, label in enumerate(headers, start=1):
            cell = ws.cell(row=4, column=col_idx, value=label)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="4F46E5")
            cell.alignment = Alignment(horizontal="left", vertical="center")

        # Data rows
        tier_fill = {
            "high": PatternFill("solid", fgColor="ECFDF5"),
            "medium": PatternFill("solid", fgColor="FFFBEB"),
            "low": PatternFill("solid", fgColor="FEF2F2"),
        }
        for offset, r in enumerate(rows, start=5):
            ws.cell(row=offset, column=1, value=r["question"]).alignment = (
                Alignment(wrap_text=True, vertical="top")
            )
            ws.cell(row=offset, column=2, value=r["answer"]).alignment = (
                Alignment(wrap_text=True, vertical="top")
            )
            tier_cell = ws.cell(
                row=offset, column=3, value=(r["confidence"] or "—").upper()
            )
            tier_cell.alignment = Alignment(horizontal="center", vertical="top")
            fill = tier_fill.get(r["confidence"])
            if fill:
                for col_idx in (1, 2, 3):
                    ws.cell(row=offset, column=col_idx).fill = fill

        # Column widths
        ws.column_dimensions[get_column_letter(1)].width = 38
        ws.column_dimensions[get_column_letter(2)].width = 56
        ws.column_dimensions[get_column_letter(3)].width = 14

        # Freeze the header
        ws.freeze_panes = "A5"

        wb.save(output_path)

    # ---- PDF (canonical printable) ----

    @staticmethod
    def _write_canonical_pdf(
        rows: list[dict[str, str]],
        output_path: str,
        *,
        title: str,
    ) -> None:
        """Render the answers as a clean PDF report.

        Uses reportlab platypus so the layout flows across pages and
        long answers wrap correctly.
        """
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
        )

        styles = getSampleStyleSheet()
        body_style = ParagraphStyle(
            name="body",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9,
            leading=12,
        )
        title_style = ParagraphStyle(
            name="title",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=16,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=4,
        )
        summary_style = ParagraphStyle(
            name="summary",
            parent=styles["BodyText"],
            fontName="Helvetica-Oblique",
            fontSize=9,
            textColor=colors.HexColor("#6B7280"),
            spaceAfter=12,
        )

        doc = SimpleDocTemplate(
            output_path,
            pagesize=LETTER,
            leftMargin=0.6 * inch,
            rightMargin=0.6 * inch,
            topMargin=0.6 * inch,
            bottomMargin=0.6 * inch,
        )
        story: list = []

        story.append(Paragraph(title, title_style))
        high = sum(1 for r in rows if r["confidence"] == "high")
        medium = sum(1 for r in rows if r["confidence"] == "medium")
        low = sum(1 for r in rows if r["confidence"] == "low")
        story.append(
            Paragraph(
                f"{len(rows)} answers · {high} high · {medium} medium · {low} low confidence",
                summary_style,
            )
        )

        if rows:
            # Build the Q&A table. Wrap each cell in a Paragraph so
            # long answers flow correctly instead of overflowing.
            data = [
                [
                    Paragraph("<b>Question</b>", body_style),
                    Paragraph("<b>Answer</b>", body_style),
                    Paragraph("<b>Conf.</b>", body_style),
                ]
            ]
            tier_color = {
                "high": colors.HexColor("#ECFDF5"),
                "medium": colors.HexColor("#FFFBEB"),
                "low": colors.HexColor("#FEF2F2"),
            }
            row_styles = []
            for idx, r in enumerate(rows, start=1):
                data.append([
                    Paragraph(_escape_pdf(r["question"]), body_style),
                    Paragraph(_escape_pdf(r["answer"]), body_style),
                    Paragraph(
                        (r["confidence"] or "—").upper(),
                        body_style,
                    ),
                ])
                bg = tier_color.get(r["confidence"])
                if bg:
                    row_styles.append(("BACKGROUND", (0, idx), (-1, idx), bg))

            table = Table(
                data,
                colWidths=[2.4 * inch, 4.0 * inch, 0.6 * inch],
                repeatRows=1,
            )
            base_styles = [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4F46E5")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D1D5DB")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
            table.setStyle(TableStyle(base_styles + row_styles))
            story.append(table)
        else:
            story.append(Paragraph("(No answers produced.)", body_style))

        story.append(Spacer(1, 0.2 * inch))
        doc.build(story)

    # -----------------------------------------------------------------------
    # Text fallback
    # -----------------------------------------------------------------------

    def _write_text_fallback(
        self, path: Path, fields: list[FilledField], output_path: str
    ) -> str:
        """For unknown formats, create a companion text file with the answers."""
        lines = ["# Tender Form — Filled Fields\n"]
        lines.append(f"Original file: {path.name}\n")
        lines.append("=" * 60 + "\n")

        for ff in fields:
            if ff.value:
                lines.append(f"{ff.name}: {ff.value}")
                if ff.confidence < 0.7:
                    lines.append(f"  ⚠ Low confidence ({ff.confidence:.0%})")
                lines.append("")

        # Write as .txt alongside the original
        txt_output = str(Path(output_path).with_suffix(".txt"))
        with open(txt_output, "w") as f:
            f.write("\n".join(lines))

        # Also copy the original
        shutil.copy2(str(path), output_path)
        print(f"Text fallback: answers written to {txt_output}")
        return output_path


# ===========================================================================
# Module-level helpers for the flat-PDF overlay path
#
# Kept at module scope (not on FormWriter) so they're trivially unit-testable
# without spinning up a FormWriter instance.  See _write_pdf_overlay_fill
# for how they compose.
# ===========================================================================


def _normalise_label(s: str) -> str:
    """Lowercase, collapse whitespace, strip trailing punctuation,
    and fold Unicode quotation marks to their ASCII equivalents.

    Used on both the FormParser-extracted field name AND the PDF word
    stream so matches survive minor punctuation differences.

    The quotation-mark fold is essential: federal PDFs almost always
    render U+2019 RIGHT SINGLE QUOTATION MARK in words like
    "Contractor's", while the FormParser-extracted field name comes
    through as ASCII U+0027 APOSTROPHE.  Without folding, the match
    misses silently.
    """
    import re

    s = s.lower()
    # Fold the common "smart quote" variants to their ASCII forms so
    # "Contractor’s" (U+2019) matches "Contractor's" (U+0027) and
    # likewise for double quotes / em-dashes used in form titles.
    s = (
        s.replace("’", "'")  # right single quotation mark
        .replace("‘", "'")  # left single quotation mark
        .replace("“", '"')  # left double quotation mark
        .replace("”", '"')  # right double quotation mark
        .replace("–", "-")  # en dash
        .replace("—", "-")  # em dash
    )
    s = re.sub(r"\s+", " ", s)
    s = s.strip()
    # Strip trailing colons, asterisks, dashes, and ellipses that
    # often appear on form labels but not in the extracted field name.
    s = s.rstrip(":*-—– ").strip()
    return s


# ---------------------------------------------------------------------------
# Geometry primitives for the rectangle-driven overlay strategy
# ---------------------------------------------------------------------------


# Tuning knobs for label clustering around each detected box.
LABEL_ROW_BAND = 5.0       # vertical tolerance for "same row" label match
LABEL_GAP_LIMIT = 18.0     # max horizontal gap inside a single label cluster
LABEL_ABOVE_REACH = 55.0   # how far ABOVE the box to scan (multi-line wrapped labels)
LABEL_ABOVE_X_SLOP = 30.0  # horizontal slop when looking for labels ABOVE the box


def _detect_input_boxes(page) -> list[dict]:
    """Detect input rectangles on one PDF page.

    Algorithm:
      1. Collect every rectangle that looks like a thin horizontal
         strip (width >= 20pt, height <= 2.5pt).  These are the
         top/bottom edges of input boxes drawn by the PDF.
      2. Pair top + bottom strips whose x-ranges match (within 3pt
         slop) and whose vertical separation is 6..30pt — the
         typical interior height of an input box.
      3. Each pair → a candidate box.
      4. Dedupe overlapping boxes: when two boxes overlap >= 70% by
         area, keep the OUTER one (larger area).  PDFs frequently
         draw each input as two concentric frames (the inner one
         being the "highlight" inside the outer border).
      5. Filter pathological detections: skip if width < 25pt or
         height < 6pt — too small to be a real input area.

    Returns list of dicts ``{x0, x1, top, bottom, width, height}``.
    """
    rects = list(getattr(page, "rects", []) or [])
    h_segs: list[dict] = []
    for r in rects:
        w = float(r["x1"]) - float(r["x0"])
        h = float(r["bottom"]) - float(r["top"])
        if w >= 20.0 and h <= 2.5:
            h_segs.append(r)
    h_segs.sort(key=lambda r: (round(float(r["top"]), 1), float(r["x0"])))

    boxes: list[dict] = []
    consumed_segs: set[int] = set()
    for i, top in enumerate(h_segs):
        if i in consumed_segs:
            continue
        for j in range(i + 1, len(h_segs)):
            if j in consumed_segs:
                continue
            bot = h_segs[j]
            if (
                abs(float(top["x0"]) - float(bot["x0"])) > 3.0
                or abs(float(top["x1"]) - float(bot["x1"])) > 3.0
            ):
                continue
            dy = float(bot["top"]) - float(top["bottom"])
            if dy < 6.0 or dy > 30.0:
                continue
            x0 = min(float(top["x0"]), float(bot["x0"]))
            x1 = max(float(top["x1"]), float(bot["x1"]))
            y_top = float(top["top"])
            y_bottom = float(bot["bottom"])
            w = x1 - x0
            h = y_bottom - y_top
            if w < 25.0 or h < 6.0:
                continue
            boxes.append({
                "x0": x0, "x1": x1,
                "top": y_top, "bottom": y_bottom,
                "width": w, "height": h,
            })
            consumed_segs.add(i)
            consumed_segs.add(j)
            break

    # Dedupe: when boxes overlap >= 70%, keep the SMALLER (inner) one.
    # Forms often draw each input as a frame-within-a-frame:
    #   outer = the table cell border
    #   inner = the actual input area inside the cell
    # We want text to land in the input area, not the cell — so we
    # take the smaller, more-deeply-nested rectangle.
    def _overlap_ratio(a: dict, b: dict) -> float:
        ix0 = max(a["x0"], b["x0"])
        ix1 = min(a["x1"], b["x1"])
        iy0 = max(a["top"], b["top"])
        iy1 = min(a["bottom"], b["bottom"])
        if ix1 <= ix0 or iy1 <= iy0:
            return 0.0
        inter = (ix1 - ix0) * (iy1 - iy0)
        a_area = a["width"] * a["height"]
        b_area = b["width"] * b["height"]
        return inter / min(a_area, b_area)

    # Sort by area ASCENDING so smaller (inner) boxes are evaluated first.
    # If a larger box overlaps a smaller one already kept, it's the outer
    # cell border — discard it.
    boxes.sort(key=lambda b: b["width"] * b["height"])
    kept: list[dict] = []
    for cand in boxes:
        if any(_overlap_ratio(cand, k) >= 0.70 for k in kept):
            continue
        kept.append(cand)

    # CRITICAL filter: drop boxes that have label TEXT inside them.
    # Tables-of-cells (like the HHS form's 4-column option grid) draw
    # each cell with the same edge pattern as a real input box.  The
    # difference is that LABEL CELLS contain printed text, INPUT
    # CELLS are empty.  By counting non-whitespace characters that
    # fall fully INSIDE the box bounds we can tell them apart.
    #
    # Tolerance: up to MAX_TEXT_CHARS_INSIDE (default 3) chars is OK,
    # so a lone "$" or "%" annotation drawn inside an input box (rare
    # but happens) doesn't get the box rejected.
    MAX_TEXT_CHARS_INSIDE = 3
    words = list(getattr(page, "extract_words", lambda **_: [])(
        x_tolerance=2, y_tolerance=2,
        keep_blank_chars=False, use_text_flow=True,
    ))

    def _chars_inside(box: dict) -> int:
        n = 0
        for w in words:
            wx0 = float(w["x0"])
            wx1 = float(w["x1"])
            wtop = float(w["top"])
            wbot = float(w["bottom"])
            # Require the word to be FULLY inside the box (with 1pt
            # slop on each edge).  A word that merely overlaps the
            # box border doesn't count — it belongs to whatever's
            # outside.
            if (
                wx0 >= box["x0"] - 1.0
                and wx1 <= box["x1"] + 1.0
                and wtop >= box["top"] - 1.0
                and wbot <= box["bottom"] + 1.0
            ):
                n += len(w["text"])
                if n > MAX_TEXT_CHARS_INSIDE:
                    return n
        return n

    input_boxes = [b for b in kept if _chars_inside(b) <= MAX_TEXT_CHARS_INSIDE]

    # Sort the surviving input boxes top-to-bottom, left-to-right
    # (matches reading order, makes the log easier to follow).
    input_boxes.sort(key=lambda b: (round(b["top"], 1), b["x0"]))
    return input_boxes


def _build_label_context(box: dict, words: list[dict]) -> str:
    """Find the label text for one box.

    Picks the label by *adjacency*, not by a flat rectangular search
    window.  Two strategies, tried in order:

    Strategy 1 — SAME-ROW cluster (the dominant case):
        1. Collect words whose vertical midpoint sits within
           LABEL_ROW_BAND of the box's midpoint AND whose right edge
           ends at or before the box's left edge.
        2. Walk those words from right to left, accumulating a
           "cluster".  Stop as soon as the horizontal gap between
           the next-leftward word and the cluster exceeds
           LABEL_GAP_LIMIT — that gap marks the boundary between
           this box's label and a different field sitting further
           to the left.
        3. The accumulated cluster IS the label.

    Strategy 2 — fallback: line(s) ABOVE the box.
        Used when no same-row label was found (e.g. forms that put
        the label on one line and the input area on the next).
        Collects words in the rows above the box that overlap its
        x-range (with LABEL_ABOVE_X_SLOP horizontal tolerance) up
        to LABEL_ABOVE_REACH above, and concatenates them in
        reading order.

    Both strategies return the text normalised (lower-case, smart-
    quotes folded, trailing punctuation stripped) so substring
    matching against field names is reliable.
    """
    box_x0 = float(box["x0"])
    box_x1 = float(box["x1"])
    box_top = float(box["top"])
    box_bot = float(box["bottom"])
    box_mid_y = (box_top + box_bot) / 2.0

    # ---- Strategy 1: same-row cluster ----
    # "Same row" = word's vertical MIDPOINT falls within the box's
    # vertical range (with LABEL_ROW_BAND of slop on each side).
    # Using midpoint-inside-box is the correct test — words that
    # merely brush the box's top/bottom edge (e.g. text from the
    # row above whose descenders dip to the box border) get
    # rejected, while words sitting in the upper or lower half of
    # the box still match cleanly.
    same_row = []
    y_lo = box_top - LABEL_ROW_BAND
    y_hi = box_bot + LABEL_ROW_BAND
    for w in words:
        w_mid_y = (float(w["top"]) + float(w["bottom"])) / 2.0
        if w_mid_y < y_lo or w_mid_y > y_hi:
            continue
        # Word must end at or before the box's left edge (with 2pt slop
        # for thin label punctuation that touches the box border).
        if float(w["x1"]) > box_x0 + 2.0:
            continue
        same_row.append(w)
    same_row.sort(key=lambda w: float(w["x0"]))

    if same_row:
        # Walk right-to-left from the box, growing a cluster of
        # adjacent words.  Stopping rule:
        #   - Once we've already incorporated a label-end token (a
        #     word ending in ":"), the cluster is "complete" — at
        #     that point any further gap > LABEL_GAP_LIMIT means we
        #     would cross into a DIFFERENT field's label, so we
        #     stop.
        #   - BEFORE we've seen a colon, we tolerate larger gaps —
        #     this handles the very common pattern of a structural
        #     "$" sign sitting alone between the label colon and
        #     the input box (e.g. "Option 1 (if applicable):    $
        #     [box]").  Without this, the cluster would terminate
        #     on just "$" and miss the actual label.
        rightward_first = list(reversed(same_row))
        cluster: list[dict] = []
        next_x = box_x0  # gap is measured against the box's left edge
        saw_colon = False
        for w in rightward_first:
            gap = next_x - float(w["x1"])
            if cluster and saw_colon and gap > LABEL_GAP_LIMIT:
                break
            cluster.insert(0, w)  # keep cluster in left-to-right order
            next_x = float(w["x0"])
            if w["text"].rstrip().endswith(":"):
                saw_colon = True
        if cluster:
            return _normalise_label(" ".join(w["text"] for w in cluster))

    # ---- Strategy 2: lines ABOVE the box ----
    above_x_min = box_x0 - LABEL_ABOVE_X_SLOP
    above_x_max = box_x1 + LABEL_ABOVE_X_SLOP
    above_y_min = box_top - LABEL_ABOVE_REACH
    above_y_max = box_top + 1.0
    above = [
        w for w in words
        if float(w["bottom"]) <= above_y_max
        and float(w["top"]) >= above_y_min
        and float(w["x1"]) >= above_x_min
        and float(w["x0"]) <= above_x_max
    ]
    above.sort(key=lambda w: (round(float(w["top"]), 1), float(w["x0"])))
    return _normalise_label(" ".join(w["text"] for w in above))


def _find_label_placement(
    *,
    label: str,
    page_words: list[list[dict]],
    page_dimensions: list[tuple[float, float]],
    consumed_words: set[tuple[int, int]],
) -> dict | None:
    """Find where to draw the answer for one label.

    Returns a dict::

        {
          "page_idx": int,
          "x": float,                # left edge of answer text (PDF coords)
          "y_top": float,             # top of the matched label rect
          "y_bottom": float,          # bottom of the matched label rect
          "font_size": float,         # clamped to 7..11 pt
          "max_width": float,         # available horizontal whitespace
        }

    Or ``None`` if no acceptable match was found on any page.

    Side effect: on a successful match, marks the matched word indices
    as consumed in the ``consumed_words`` set so later (shorter) labels
    cannot match into the same physical phrase.

    Matching algorithm:
        1. Tokenise the label into N words (normalised).
        2. Slide a window of N consecutive words across each page's
           word list.  Skip windows that overlap any consumed words.
        3. Accept windows where the normalised concatenated text
           either EXACTLY equals the target, or matches and the
           immediately-following PDF word would clearly start a new
           field (begins with a colon, a digit, a dollar sign, or
           is itself a short label-like token).  This guards against
           short labels like "Title" matching mid-phrase inside
           "Title of Acquisition".
        4. The label's anchor is the END of the rightmost matched
           word; the answer's start x = that x1 + 4pt padding.
        5. Available width is computed as the distance to the next
           word on the same vertical band (within ±3pt of the
           label's centerline).  If there is no such word, it falls
           through to (page_width - x - 40pt) which leaves a 40pt
           right-margin breathing room.
    """
    target_norm = _normalise_label(label)
    if not target_norm:
        return None
    target_tokens = target_norm.split()
    if not target_tokens:
        return None
    n = len(target_tokens)

    # Words that, when they immediately follow the matched window,
    # tell us we matched mid-phrase rather than at a label boundary.
    # If the next word is one of these, reject the match (unless the
    # window is already an exact match for the whole target).
    CONNECTOR_WORDS = {
        "of", "and", "or", "for", "in", "to", "with", "the", "by",
        "&", "/", "-",
    }

    for page_idx, words in enumerate(page_words):
        if len(words) < n:
            continue
        for i in range(len(words) - n + 1):
            # Reject any candidate window that overlaps already-consumed
            # words (prevents short labels claiming words inside longer
            # labels that were placed earlier this pass).
            if any((page_idx, idx) in consumed_words for idx in range(i, i + n)):
                continue

            window = words[i : i + n]
            window_text = _normalise_label(" ".join(w["text"] for w in window))

            # Require either exact match (covers the typical case) OR
            # a strict-prefix match plus a label-boundary signal on
            # the very next word.
            exact = window_text == target_norm
            prefix = (
                not exact
                and (
                    window_text.startswith(target_norm + " ")
                    or window_text.startswith(target_norm)
                )
            )
            if not exact and not prefix:
                continue

            # Anti-substring guard: if not exact, peek at the next
            # word in the page's word list.  If it's a connector
            # word like "of"/"and", we're likely matching INSIDE a
            # longer phrase — reject.
            if not exact and i + n < len(words):
                next_text = _normalise_label(words[i + n]["text"])
                # Strip any trailing colon — a colon means label
                # boundary, which is fine.
                if next_text and not next_text.startswith(":"):
                    next_first_token = next_text.split()[0] if next_text else ""
                    if next_first_token in CONNECTOR_WORDS:
                        continue

            last = window[n - 1]
            label_y_top = float(last["top"])
            label_y_bottom = float(last["bottom"])
            label_mid_y = (label_y_top + label_y_bottom) / 2.0

            # ---------------- Annotation skip ----------------
            # The bare label often has trailing annotations that
            # aren't part of the answer area:
            #   "Base Period (if there are options): $ ______"
            #   "Total Modification Amount: (if applicable) $ __"
            # FormParser extracts the field name as just
            # "Base Period" / "Total Modification Amount" — but if we
            # place answer text immediately after the last matched
            # word, we'd overlap the "(if there are options):" or "$"
            # annotation glyphs.  Advance past the annotation cluster
            # until we hit either:
            #   (a) end of the same line, or
            #   (b) a SECOND colon outside parens (= new label
            #       starting after the answer area, e.g.
            #       "OPDIV/Division/Branch (including location): ___
            #        Email: ___" — we must stop BEFORE "Email:"), or
            #   (c) a word that is clearly NOT an annotation —
            #       has no leading `(`, isn't `:` / `$`, and isn't
            #       inside a `(...)` group.
            # Cap to 8 words of look-ahead so a runaway scan can't
            # cross into the next field on a tightly packed form.
            advance_end = i + n - 1  # index of `last` matched word
            paren_depth = 0
            # If the matched window itself ended with `:` (e.g.
            # "...(including location):" matched in one go), count
            # that as 1 so the very next outside-paren colon during
            # the skip — which would be a new label — triggers stop.
            last_text = window[-1]["text"].rstrip()
            colon_outside_paren = 1 if last_text.endswith(":") else 0
            for j in range(i + n, min(i + n + 8, len(words))):
                w = words[j]
                wt = w["text"].strip()
                w_mid = (float(w["top"]) + float(w["bottom"])) / 2.0
                # Stop if we've left the label's vertical band
                if abs(w_mid - label_mid_y) > 3.0:
                    break

                # Compute the paren-depth this word WOULD produce —
                # we look at it before deciding whether to consume.
                new_paren_depth = paren_depth + wt.count("(") - wt.count(")")

                # SECOND outside-paren colon = next label on the
                # same line.  Stop BEFORE consuming this word.
                if new_paren_depth <= 0 and wt.endswith(":"):
                    if colon_outside_paren >= 1:
                        break
                    colon_outside_paren += 1

                paren_depth = new_paren_depth

                is_annotation = (
                    wt in (":", "$", "$:", "*", "(", ")")
                    or wt.startswith("(")
                    or wt.endswith(":")
                    or wt.endswith("$")
                    or paren_depth > 0
                )
                if is_annotation:
                    advance_end = j
                    continue
                # First non-annotation word — stop, don't consume it
                break

            anchor_word = words[advance_end]
            x_start = float(anchor_word["x1"]) + 4.0
            y_top = float(anchor_word["top"])
            y_bottom = float(anchor_word["bottom"])

            page_w, _page_h = page_dimensions[page_idx]
            right_limit = page_w - 40.0

            for w in words:
                if w is anchor_word:
                    continue
                wx0 = float(w["x0"])
                if wx0 <= x_start:
                    continue
                w_top = float(w["top"])
                w_bot = float(w["bottom"])
                if (
                    abs(((w_top + w_bot) / 2.0) - label_mid_y) <= 3.0
                    and wx0 < right_limit
                ):
                    right_limit = wx0 - 4.0
            max_width = max(20.0, right_limit - x_start)

            label_h = max(1.0, y_bottom - y_top)
            font_size = max(7.0, min(11.0, label_h - 1.0))

            # Side-effect: claim every word index in this match AND
            # every annotation word we walked past, so a later label
            # can't grab them.
            for idx in range(i, advance_end + 1):
                consumed_words.add((page_idx, idx))

            return {
                "page_idx": page_idx,
                "x": x_start,
                "y_top": y_top,
                "y_bottom": y_bottom,
                "font_size": font_size,
                "max_width": max_width,
            }

    return None


def _truncate_to_width(text: str, max_width: float, font_size: float) -> str:
    """Truncate ``text`` so it fits in ``max_width`` at ``font_size``.

    Helvetica isn't a fixed-width font — avg glyph advance is roughly
    0.50 × font_size (digits) to 0.55 × font_size (letters).  We use
    0.52 as a midpoint and add "…" when truncating.
    """
    if not text:
        return ""
    avg_char_width = font_size * 0.52
    if avg_char_width <= 0:
        return text
    max_chars = int(max_width / avg_char_width)
    if max_chars >= len(text):
        return text
    if max_chars < 4:
        return text[:max(0, max_chars)]
    return text[: max_chars - 1] + "…"
