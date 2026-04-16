"""
Document Parser — Extracts structured text from PDF, DOCX, TXT, and MD files.

WHY THIS EXISTS:
The RAG pipeline needs raw text to chunk and embed. But documents come in different
formats, each requiring a different extraction strategy. This module provides a single
`parse()` method that auto-detects the file type and returns a list of ParsedPage
objects — one per page (for PDFs) or one per logical section (for DOCX/TXT).

Each ParsedPage carries metadata (source file, page number, section heading) that
flows through chunking and into the vector database, so when the RAG pipeline
retrieves a chunk during tender drafting, it can cite exactly where the information
came from.

KEY DESIGN DECISIONS:
- pdfplumber over PyPDF2: Better text extraction quality, especially for tables
  and multi-column layouts common in government tender documents.
- python-docx for DOCX: Gives us paragraph-level access with style info (Heading 1,
  Heading 2, Normal), which lets us group text by section — much better chunks.
- One ParsedPage per PDF page: Government tenders often reference "see page X",
  so page-level granularity matters for source attribution.
- One ParsedPage per heading section for DOCX: Company docs are usually organised
  by headings (e.g., "Company Overview", "Certifications", "Pricing"), and those
  sections map naturally to tender response sections.
"""

from __future__ import annotations

import hashlib
import logging
from dataclasses import dataclass, field
from pathlib import Path

import structlog

logger = structlog.get_logger(__name__)

# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class ParsedPage:
    """One logical unit of extracted text with its source metadata.

    For PDFs: one ParsedPage = one physical page.
    For DOCX: one ParsedPage = one heading section (or the whole doc if no headings).
    For TXT/MD: one ParsedPage = the entire file content.

    Attributes:
        source_file: Original filename (not full path — for portability).
        page_number: 1-indexed page number (PDFs) or section index (DOCX).
        section_heading: The heading text for this section (DOCX only), or None.
        text: The extracted text content.
        char_count: Number of characters in `text`.
        content_hash: SHA-256 hash of `text` — used to detect duplicate ingestion.
    """
    source_file: str
    page_number: int
    section_heading: str | None
    text: str
    char_count: int = field(init=False)
    content_hash: str = field(init=False)

    def __post_init__(self) -> None:
        self.char_count = len(self.text)
        self.content_hash = hashlib.sha256(self.text.encode("utf-8")).hexdigest()


# ---------------------------------------------------------------------------
# Supported file extensions
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS: set[str] = {".pdf", ".docx", ".txt", ".md"}


# ---------------------------------------------------------------------------
# Parser class
# ---------------------------------------------------------------------------

class DocumentParser:
    """Extracts text from documents in various formats.

    Usage:
        parser = DocumentParser()
        pages = parser.parse("data/knowledge_base/company_overview.pdf")
        for page in pages:
            print(f"Page {page.page_number}: {page.char_count} chars")
    """

    def parse(self, file_path: str | Path) -> list[ParsedPage]:
        """Parse a document and return a list of ParsedPage objects.

        Args:
            file_path: Path to the document file.

        Returns:
            List of ParsedPage objects, one per logical section.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file extension is not supported.
            RuntimeError: If parsing fails for any reason.
        """
        path = Path(file_path)

        extension = path.suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: '{extension}'. "
                f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        logger.info(
            "parsing_document",
            file=path.name,
            extension=extension,
            size_bytes=path.stat().st_size,
        )

        try:
            if extension == ".pdf":
                pages = self._parse_pdf(path)
            elif extension == ".docx":
                pages = self._parse_docx(path)
            else:
                # .txt and .md — plain text
                pages = self._parse_text(path)
        except Exception as exc:
            logger.error("parsing_failed", file=path.name, error=str(exc))
            raise RuntimeError(f"Failed to parse {path.name}: {exc}") from exc

        # Filter out empty pages (some PDFs have blank pages)
        pages = [p for p in pages if p.text.strip()]

        logger.info(
            "parsing_complete",
            file=path.name,
            pages_extracted=len(pages),
            total_chars=sum(p.char_count for p in pages),
        )

        return pages

    def parse_many(self, file_paths: list[str | Path]) -> list[ParsedPage]:
        """Parse multiple documents and return all pages combined.

        Skips files that fail to parse (logs the error) rather than
        stopping the entire batch — one bad file shouldn't block
        ingestion of 50 good ones.

        Args:
            file_paths: List of file paths to parse.

        Returns:
            Combined list of ParsedPage objects from all successfully parsed files.
        """
        all_pages: list[ParsedPage] = []
        success_count = 0
        fail_count = 0

        for fp in file_paths:
            try:
                pages = self.parse(fp)
                all_pages.extend(pages)
                success_count += 1
            except (FileNotFoundError, ValueError, RuntimeError) as exc:
                logger.warning("skipping_file", file=str(fp), reason=str(exc))
                fail_count += 1

        logger.info(
            "batch_parsing_complete",
            files_succeeded=success_count,
            files_failed=fail_count,
            total_pages=len(all_pages),
        )

        return all_pages

    # ------------------------------------------------------------------
    # Private extraction methods
    # ------------------------------------------------------------------

    def _parse_pdf(self, path: Path) -> list[ParsedPage]:
        """Extract text from a PDF file, one ParsedPage per physical page.

        Uses pdfplumber which handles multi-column layouts, tables, and
        embedded fonts better than PyPDF2. Each page's text is extracted
        with layout preservation (keeps reading order correct).
        """
        import pdfplumber

        pages: list[ParsedPage] = []

        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                pages.append(
                    ParsedPage(
                        source_file=path.name,
                        page_number=i,
                        section_heading=None,
                        text=text,
                    )
                )

        return pages

    def _parse_docx(self, path: Path) -> list[ParsedPage]:
        """Extract text from a DOCX file, grouped by heading sections.

        HOW SECTION GROUPING WORKS:
        We iterate through every paragraph in the document. When we hit a
        paragraph styled as a heading (Heading 1, Heading 2, etc.), we
        start a new section. All paragraphs that follow — until the next
        heading — belong to that section. This gives us semantically
        meaningful chunks like "Company Overview" or "Certifications".

        If the document has NO headings at all, the entire content becomes
        a single ParsedPage (common for simple docs like team bios).

        WHY THIS MATTERS:
        When the RAG pipeline retrieves a chunk about "ISO 27001 certification",
        knowing it came from the "Certifications" section of "company_profile.docx"
        helps the drafting LLM place it in the right tender section.
        """
        import docx

        doc = docx.Document(str(path))

        sections: list[ParsedPage] = []
        current_heading: str | None = None
        current_text_parts: list[str] = []
        section_index = 1

        def _flush_section() -> None:
            """Save the current accumulated text as a ParsedPage."""
            nonlocal section_index
            text = "\n".join(current_text_parts).strip()
            if text:
                sections.append(
                    ParsedPage(
                        source_file=path.name,
                        page_number=section_index,
                        section_heading=current_heading,
                        text=text,
                    )
                )
                section_index += 1

        for paragraph in doc.paragraphs:
            style_name = (paragraph.style.name or "").lower()

            # Detect headings — python-docx uses style names like
            # "Heading 1", "Heading 2", etc.
            is_heading = style_name.startswith("heading")

            if is_heading:
                # Save whatever we've accumulated so far
                _flush_section()
                # Start a new section
                current_heading = paragraph.text.strip() or None
                current_text_parts = []
                # Include the heading text itself in the section body
                # so the chunk contains the heading for context
                if current_heading:
                    current_text_parts.append(current_heading)
            else:
                text = paragraph.text.strip()
                if text:
                    current_text_parts.append(text)

        # Don't forget the last section
        _flush_section()

        # If no sections were created (no headings in the doc), treat
        # the entire document as one page
        if not sections:
            full_text = "\n".join(
                p.text.strip() for p in doc.paragraphs if p.text.strip()
            )
            if full_text:
                sections.append(
                    ParsedPage(
                        source_file=path.name,
                        page_number=1,
                        section_heading=None,
                        text=full_text,
                    )
                )

        return sections

    def _parse_text(self, path: Path) -> list[ParsedPage]:
        """Extract text from a plain text or Markdown file.

        Simple — read the entire file as one ParsedPage. We don't try
        to split by Markdown headings here because the chunker handles
        that more consistently across all document types.
        """
        text = path.read_text(encoding="utf-8", errors="replace")

        return [
            ParsedPage(
                source_file=path.name,
                page_number=1,
                section_heading=None,
                text=text,
            )
        ]